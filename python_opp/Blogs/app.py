from flask import Flask, render_template, Response, request, redirect
from pymysql import connect, cursors, Error
from datetime import datetime  
from docx import Document
from docx.shared import Inches

config = {
    'host': 'localhost',
    'user': 'root',
    'password': '',
    'database': 'blogs',
    }
cnx = connect(**config)

#from app import app

app = Flask(__name__)


# Chỉ định URL kích hoạt hàm homepage
@app.route('/index', methods = ["POST"])
def postpage():
    title = request.form["title"]
    content = request.form["content"]
    cur = cnx.cursor()
    sql="INSERT INTO list_blog (title, content, create_date) VALUES (%s, %s,%s)"
    value=(title,content,datetime.now())
    try:
        cur.execute(sql,value)
        cnx.commit()
    except:
        cnx.rollback()
    #cnx.close()
    return redirect("/index", code=302)

@app.route('/index', methods = ["GET"])
def getpage():
    cur = cnx.cursor()
    sql="SELECT * FROM list_blog"
    cur.execute(sql)
    return render_template("index.html",cur=cur)
@app.route('/editpost/<id>', methods = ["GET"])
def detailpost(id):
    cur = cnx.cursor()
    sql="SELECT * FROM list_blog WHERE id=" +str(id)
    cur.execute(sql)
    for i in cur:
        r1 = i[1] # r1 là title của bài viết
        r2 = i[2] #r2 là content của bài viết
        r0 = i[0] #r0 là id của bài viết
    return render_template("editpost.html",r1=r1, r2=r2, r0=r0)
@app.route('/editpost/<id>', methods = ["POST"])
def editpost(id):
    cur = cnx.cursor()
    title = request.form["title"]
    content = request.form["content"]
    sql=f"UPDATE list_blog SET title = '{title}', content ='{content}' WHERE id="+str(id)
    cur.execute(sql)
    cnx.commit()
    return redirect("/index", code=302)
    
@app.route('/resignationletter', methods = ["GET"])
def resignation():
    return render_template("resignationletter.html")

@app.route('/about', methods = ["GET"])
def about():
    return render_template("about.html")

@app.route('/resignationletter', methods = ["POST"])
def resignationletter():
    cur = cnx.cursor()
    fullname = request.form["fullname"]
    reason = request.form["reason"]
    document = Document()
    document.add_paragraph(fullname)
    document.add_paragraph(reason)
    document.save('static/demo.docx')
    return redirect("/resignationletter", code=302)
   

if __name__ == "__main__":
    # Chạy Flask app
    # Bật debug để restart server mỗi khi có thay đổi trong mã
    app.run(debug=True)
